# src/document_processor/loader.py
"""
Document loaders for various file types.
"""

from pathlib import Path
from typing import List
import pypdf
import docx

class Document:
    """Simple document container"""
    def __init__(self, content: str, metadata: dict = None):
        self.content = content
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document(chars={len(self.content)}, metadata={self.metadata})"

class DocumentLoader:
    """Load documents from various file formats"""
    
    @staticmethod
    def load_txt(file_path: str) -> Document:
        """Load text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return Document(
            content=content,
            metadata={
                'source': file_path,
                'type': 'txt'
            }
        )
    
    @staticmethod
    def load_pdf(file_path: str) -> Document:
        """Load PDF file. Builds page_map so callers can map char positions → page numbers."""
        SEP = '\n\n'
        page_texts = []
        page_map = []   # list of (page_num_1based, char_start, char_end)
        pos = 0

        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                text = pdf_reader.pages[page_num].extract_text() or ""
                page_map.append((page_num + 1, pos, pos + len(text)))
                page_texts.append(text)
                pos += len(text) + len(SEP)  # account for the separator we'll join with

        return Document(
            content=SEP.join(page_texts),
            metadata={
                'source': file_path,
                'type': 'pdf',
                'pages': num_pages,
                'page_map': page_map,   # [(page, start_char, end_char), ...]
            }
        )
    
    @staticmethod
    def load_docx(file_path: str) -> Document:
        """Load Word document"""
        doc = docx.Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        return Document(
            content='\n\n'.join(content),
            metadata={
                'source': file_path,
                'type': 'docx',
                'paragraphs': len(content)
            }
        )
    
    @staticmethod
    def load(file_path: str) -> Document:
        """
        Auto-detect file type and load.
        
        Args:
            file_path: Path to file
        
        Returns:
            Document object
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.txt':
            return DocumentLoader.load_txt(file_path)
        elif extension == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif extension == '.docx':
            return DocumentLoader.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")